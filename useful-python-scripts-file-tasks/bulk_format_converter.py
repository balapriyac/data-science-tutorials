"""
Bulk Format Converter
---------------------
Converts files across multiple formats in batch — images, audio,
and documents. Set the input folder, target format, and optional
quality settings. Converted files go to a dedicated output folder.
Skips unsupported or already-matching formats and logs any failures.

Usage:
    python bulk_format_converter.py /path/to/files --format webp
    python bulk_format_converter.py /path/to/files --format mp3 --quality 192
    python bulk_format_converter.py /path/to/files --format jpg --quality 85
    python bulk_format_converter.py /path/to/files --format txt --dry-run

Dependencies:
    pip install Pillow pydub python-docx
    For audio: also requires ffmpeg installed on your system
        https://ffmpeg.org/download.html
"""

import os
import sys
import argparse
from pathlib import Path
from datetime import datetime


# ---------------------------------------------------------------------------
# Format Maps
# ---------------------------------------------------------------------------

# Supported input extensions grouped by type
IMAGE_FORMATS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff", ".tif", ".webp", ".ico", ".ppm"}
AUDIO_FORMATS = {".wav", ".mp3", ".ogg", ".flac", ".aac", ".wma", ".m4a"}
DOC_FORMATS = {".docx", ".txt"}

# Valid output formats per category
VALID_IMAGE_OUTPUTS = {"png", "jpg", "jpeg", "bmp", "gif", "tiff", "webp"}
VALID_AUDIO_OUTPUTS = {"mp3", "wav", "ogg", "flac", "aac"}
VALID_DOC_OUTPUTS = {"txt", "docx"}

# Default quality settings per output format
DEFAULT_QUALITY = {
    "jpg":  85,
    "jpeg": 85,
    "png":  None,       # PNG is lossless — quality setting is ignored
    "webp": 80,
    "mp3":  192,        # bitrate in kbps
    "ogg":  192,
    "aac":  192,
    "wav":  None,       # lossless
    "flac": None,       # lossless
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def detect_file_type(ext: str) -> str | None:
    """Determine whether a file extension belongs to images, audio, or docs."""
    ext = ext.lower()
    if ext in IMAGE_FORMATS:
        return "image"
    if ext in AUDIO_FORMATS:
        return "audio"
    if ext in DOC_FORMATS:
        return "doc"
    return None


def get_valid_outputs(file_type: str) -> set[str]:
    """Return valid output formats for a given file type."""
    return {
        "image": VALID_IMAGE_OUTPUTS,
        "audio": VALID_AUDIO_OUTPUTS,
        "doc":   VALID_DOC_OUTPUTS,
    }.get(file_type, set())


def format_size(size_bytes: int) -> str:
    """Human-readable file size."""
    for unit in ("B", "KB", "MB", "GB", "TB"):
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} PB"


# ---------------------------------------------------------------------------
# Converters
# ---------------------------------------------------------------------------

def convert_image(input_path: Path, output_path: Path, quality: int | None) -> None:
    """Convert an image file using Pillow."""
    from PIL import Image

    img = Image.open(input_path)

    # Convert to RGB if saving as JPEG (no alpha channel support)
    if output_path.suffix.lower() in (".jpg", ".jpeg") and img.mode in ("RGBA", "P", "LA"):
        img = img.convert("RGB")

    # Convert to RGBA if saving as PNG with transparency
    if output_path.suffix.lower() == ".png" and img.mode == "P":
        img = img.convert("RGBA")

    save_kwargs = {}
    if quality and output_path.suffix.lower() in (".jpg", ".jpeg", ".webp"):
        save_kwargs["quality"] = quality

    img.save(output_path, **save_kwargs)


def convert_audio(input_path: Path, output_path: Path, bitrate: int | None) -> None:
    """Convert an audio file using pydub (requires ffmpeg)."""
    from pydub import AudioSegment

    audio = AudioSegment.from_file(str(input_path))

    export_kwargs = {"format": output_path.suffix.lstrip(".")}

    if bitrate and output_path.suffix.lower() in (".mp3", ".ogg", ".aac"):
        export_kwargs["bitrate"] = f"{bitrate}k"

    audio.export(str(output_path), **export_kwargs)


def convert_doc(input_path: Path, output_path: Path) -> None:
    """Convert between docx and txt formats."""
    input_ext = input_path.suffix.lower()
    output_ext = output_path.suffix.lower()

    if input_ext == ".docx" and output_ext == ".txt":
        # DOCX -> TXT
        from docx import Document
        doc = Document(str(input_path))
        text = "\n".join([para.text for para in doc.paragraphs])
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)

    elif input_ext == ".txt" and output_ext == ".docx":
        # TXT -> DOCX
        from docx import Document
        with open(input_path, "r", encoding="utf-8") as f:
            text = f.read()
        doc = Document()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(str(output_path))

    else:
        raise ValueError(f"Unsupported doc conversion: {input_ext} → {output_ext}")


# ---------------------------------------------------------------------------
# Core Logic
# ---------------------------------------------------------------------------

class ConversionResult:
    """Tracks outcome of a single file conversion."""
    def __init__(self, source: str, target: str, status: str, error: str = ""):
        self.source = source
        self.target = target
        self.status = status      # "success", "skipped", "failed"
        self.error = error


def run_conversion(
    input_dir: Path,
    output_dir: Path,
    target_format: str,
    quality: int | None,
    dry_run: bool = False,
    recursive: bool = True,
) -> list[ConversionResult]:
    """
    Scan input_dir for convertible files and convert them to target_format.
    """
    results = []
    target_ext = f".{target_format.lower()}"

    # Gather files
    glob_pattern = "**/*" if recursive else "*"
    all_files = [f for f in input_dir.glob(glob_pattern) if f.is_file()]

    for file_path in sorted(all_files):
        ext = file_path.suffix.lower()
        file_type = detect_file_type(ext)

        # Skip files we don't recognize
        if file_type is None:
            continue

        # Skip files already in the target format
        if ext == target_ext:
            results.append(ConversionResult(str(file_path), "", "skipped", "Already in target format"))
            continue

        # Validate the conversion is supported
        valid_outputs = get_valid_outputs(file_type)
        if target_format.lower() not in valid_outputs:
            results.append(ConversionResult(
                str(file_path), "", "skipped",
                f"Cannot convert {file_type} to .{target_format}"
            ))
            continue

        # Build output path
        output_path = output_dir / f"{file_path.stem}{target_ext}"

        # Handle name conflicts
        counter = 1
        while output_path.exists():
            output_path = output_dir / f"{file_path.stem}_{counter}{target_ext}"
            counter += 1

        if dry_run:
            print(f"   [DRY RUN] {file_path.name} → {output_path.name}")
            results.append(ConversionResult(str(file_path), str(output_path), "would convert"))
            continue

        # Perform conversion
        try:
            print(f"   Converting: {file_path.name} → {output_path.name}", end=" ", flush=True)

            if file_type == "image":
                convert_image(file_path, output_path, quality)
            elif file_type == "audio":
                convert_audio(file_path, output_path, quality)
            elif file_type == "doc":
                convert_doc(file_path, output_path)

            print("✓")
            results.append(ConversionResult(str(file_path), str(output_path), "success"))

        except Exception as e:
            print(f"✗ ({e})")
            results.append(ConversionResult(str(file_path), str(output_path), "failed", str(e)))

    return results


def write_log(results: list[ConversionResult], log_path: Path) -> None:
    """Write a conversion log file."""
    with open(log_path, "w", encoding="utf-8") as f:
        f.write(f"Conversion Log — {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("=" * 70 + "\n\n")

        for r in results:
            f.write(f"[{r.status.upper()}]\n")
            f.write(f"  Source: {r.source}\n")
            if r.target:
                f.write(f"  Output: {r.target}\n")
            if r.error:
                f.write(f"  Note:   {r.error}\n")
            f.write("\n")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Batch convert files between image, audio, and document formats."
    )
    parser.add_argument(
        "input", type=str,
        help="Input directory containing files to convert"
    )
    parser.add_argument(
        "--format", type=str, required=True,
        help="Target format (e.g. webp, mp3, jpg, txt, png, wav, docx)"
    )
    parser.add_argument(
        "--quality", type=int, default=None,
        help="Quality/bitrate setting (e.g. 85 for JPEG, 192 for MP3). Uses defaults if not specified."
    )
    parser.add_argument(
        "--output", type=str, default=None,
        help="Output directory (default: 'converted' folder inside input directory)"
    )
    parser.add_argument(
        "--dry-run", action="store_true",
        help="Show what would be converted without converting anything"
    )
    parser.add_argument(
        "--no-recursive", action="store_true",
        help="Only process files in the top-level directory (default: scan subdirectories too)"
    )
    args = parser.parse_args()

    input_dir = Path(args.input).resolve()
    if not input_dir.exists():
        print(f"❌ Input directory not found: {input_dir}")
        sys.exit(1)

    target_format = args.format.lower().lstrip(".")

    # Set up output
    if args.output:
        output_dir = Path(args.output).resolve()
    else:
        output_dir = input_dir / "converted"

    if not args.dry_run:
        output_dir.mkdir(parents=True, exist_ok=True)

    # Determine quality
    quality = args.quality if args.quality else DEFAULT_QUALITY.get(target_format)

    # Info
    print(f"\n📂 Input:   {input_dir}")
    print(f"📁 Output:  {output_dir}")
    print(f"🎯 Target:  .{target_format}")
    if quality:
        print(f"⚙️  Quality: {quality}")
    if args.dry_run:
        print("🛑 DRY RUN — no files will be converted")
    print()

    # Run
    results = run_conversion(
        input_dir, output_dir, target_format, quality,
        dry_run=args.dry_run, recursive=not args.no_recursive
    )

    # Summary
    success  = sum(1 for r in results if r.status == "success")
    skipped  = sum(1 for r in results if r.status in ("skipped", "would convert"))
    failed   = sum(1 for r in results if r.status == "failed")

    print(f"\n{'=' * 50}")
    print(f"  ✅ Conversion complete")
    print(f"  Converted:  {success}")
    print(f"  Skipped:    {skipped}")
    if failed:
        print(f"  Failed:     {failed}")
    print(f"{'=' * 50}")

    # Write log
    if not args.dry_run:
        log_path = output_dir / "conversion_log.txt"
        write_log(results, log_path)
        print(f"\n📋 Log saved: {log_path}")


if __name__ == "__main__":
    main()


